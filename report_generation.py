from docx import Document
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import datetime
import pandas as pd
from docx.shared import RGBColor
from docx.enum.style import *


def automatic_report_generate(customer_name,pb_name,pb_comment,els_df,target_bond):
    document = Document('./data/템플릿.docx')
    today_date = datetime.datetime.now().strftime("%Y.%m.%d")

    paragraph3 = document.add_paragraph()

    run = paragraph3.add_run('{0} 고객님 귀하\n'.format(customer_name))
    run.bold = True
    run.underline = True
    paragraph3.add_run(today_date).bold = True

    paragraph3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    head = document.add_heading("", 0)
    head.add_run(pb_comment).bold=True
    head.runs[0].font.size=Pt(30)

    document.add_paragraph("_________________________________________________________________________________________________")
    head = document.add_heading('', level = 0)
    head.add_run('증시').bold=True

    table = document.add_table(rows = 2, cols = 7)
    table.style = document.styles['Table Grid']
    first_row = table.rows[0].cells
    first_row[0].text = '코스피'
    first_row[1].text = '코스닥'
    first_row[2].text = 'S&P500'
    first_row[3].text = '나스닥'
    first_row[4].text = '국채3년'
    first_row[5].text = '국채10년'
    first_row[6].text = '유가'

    second_row = table.rows[1].cells
    domestic = pd.read_csv("./data/국내지수.csv")
    nondomestic = pd.read_csv("./data/해외지수.csv")
    wti = pd.read_csv("./data/WTI.csv")
    domestic_bond = pd.read_csv("./data/국채데이터.csv")

    import numpy as np

    domestic['코스피변화율'] = np.round(domestic['Kospi'].pct_change()*100,2)
    domestic['코스닥변화율'] = np.round(domestic['Kosdaq'].pct_change()*100,2)
    nondomestic['나스닥변화율'] = np.round(nondomestic['Nasdaq'].pct_change()*100,2)
    nondomestic['S&P변화율'] = np.round(nondomestic['S&P500'].pct_change()*100,2)

    wti['변화율'] = np.round(wti['Close'].pct_change()*100,2)
    domestic_bond_3year = domestic_bond[domestic_bond['종목명'] == "국채3년"]
    domestic_bond_10year = domestic_bond[domestic_bond['종목명'] == "국채10년"]

    domestic['Kospi'] = np.round(domestic['Kospi'],2)
    second_row[0].text = str(domestic['코스피변화율'].tail(1).values[0]) +"%"

    def sign_define(x,close,locate,bond=False):
        close = np.round(close,2)
        if bond:
            if x>0:
                locate.text = str(close) + "\n(+" + str(x) + ")"
            elif x==0:
                locate.text = str(close) + "\n(" + str(x) + ")"
            else:
                locate.text = str(close) + "\n(" + str(x) + ")"
        else:
            if x>0:
                locate.text = str(close) + "\n(+" + str(x) + "%)"
            elif x==0:
                locate.text = str(close) + "\n(" + str(x) + ")"
            else:
                locate.text = str(close) + "\n(" + str(x) + "%)"

    sign_define(domestic['코스피변화율'].tail(1).values[0],domestic['Kospi'].tail(1).values[0] ,second_row[0])
    sign_define(domestic['코스닥변화율'].tail(1).values[0],domestic['Kosdaq'].tail(1).values[0] ,second_row[1])
    sign_define(nondomestic['S&P변화율'].tail(1).values[0],nondomestic['S&P500'].tail(1).values[0],second_row[2])
    sign_define(nondomestic['나스닥변화율'].tail(1).values[0],nondomestic['Nasdaq'].tail(1).values[0],second_row[3])
    sign_define(domestic_bond_3year['대비'].tail(1).values[0],domestic_bond_3year['수익률'].tail(1).values[0],second_row[4],True)
    sign_define(domestic_bond_10year['대비'].tail(1).values[0],domestic_bond_10year['수익률'].tail(1).values[0],second_row[5],True)
    sign_define(wti['변화율'].tail(1).values[0],wti['Close'].tail(1).values[0],second_row[6])

    document.add_paragraph("_________________________________________________________________________________________________")

    head = document.add_heading('', level = 0)
    head.add_run("금주 추천 금융 상품 ({0} ~ {1})".format(today_date,(datetime.datetime.today()+datetime.timedelta(days=7)).strftime("%Y.%m.%d"))).bold=True

    head = document.add_heading('', level = 0)
    head.add_run('1.회사채').bold=True

    table = document.add_table(target_bond.shape[0]+1, target_bond.shape[1],style = document.styles['Table Grid'])

    target_columns = target_bond.columns

    for r in range(len(table.rows)):
        row = table.rows[r]
        for c in range(len(row.cells)):
            if r == 0:
                cell = row.cells[c]
                cell.text = str(target_bond.columns[c])
            else:
                cell = row.cells[c]
                cell.text = str(target_bond.iloc[r-1][target_columns[c]])

    document.add_paragraph('\n')

    head = document.add_heading('', level = 0)
    head.add_run('2.ELS').bold=True

    table = document.add_table(els_df.shape[0]+1, els_df.shape[1],style = document.styles['Table Grid'])
    table.autofit = True
    target_columns = els_df.columns

    for r in range(len(table.rows)):
        row = table.rows[r]
        for c in range(len(row.cells)):
            if r == 0:
                cell = row.cells[c]
                cell.text = str(els_df.columns[c])
            else:
                cell = row.cells[c]
                cell.text = str(els_df.iloc[r-1][target_columns[c]])

    document.add_paragraph('\n')
    document.add_paragraph("_________________________________________________________________________________________________")

    head = document.add_heading('', level = 0)
    head.add_run('보유종목 Report').bold=True

    table = document.add_table(rows = 1,cols= 2)
    table.autofit = True
    paragraph = table.rows[0].cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("./Image/Default2"+ "포트폴리오.png", width=Inches(4.3),height=Inches(3.5))
    stock_price = pd.read_csv("./data/stock_price.csv")

    stock_close_price = []
    stock_pct_change = []
    target_stock = list(stock_price.columns[1:].values)

    for i in stock_price.columns[1:]:
        stock_pct_change.append(np.round(stock_price[i].pct_change().tail(1).values[0]*100,2))
        stock_close_price.append(int(stock_price[i].tail(1).values[0]))

    target_string = "\n\n\n\n"
    for idx,stock_name in enumerate(target_stock):
        target_string += stock_name + "\n\t"
        target_string += str(stock_close_price[idx]) + "원 (" + str(stock_pct_change[idx]) + "%)\n\n"

    table.rows[0].cells[0].text = target_string

    paragraph3 = document.add_paragraph('\n')
    paragraph3.add_run('{0}, 여의도 영업부/113240@koreainvestment.com'.format(pb_name)).bold = True

    paragraph3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save("./Generated_Report/{0}고객님 리포트.docx".format(customer_name))

